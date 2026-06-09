from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO

import fitz
from docx import Document


def _read_bytes(file: BinaryIO) -> bytes:
    if hasattr(file, "getvalue"):
        return file.getvalue()
    current_position = file.tell() if hasattr(file, "tell") else None
    if hasattr(file, "seek"):
        file.seek(0)
    data = file.read()
    if current_position is not None and hasattr(file, "seek"):
        file.seek(current_position)
    return data


def read_pdf(file: BinaryIO) -> str:
    """Read text from all PDF pages using PyMuPDF."""
    data = _read_bytes(file)
    chunks: list[str] = []

    with fitz.open(stream=data, filetype="pdf") as document:
        for index, page in enumerate(document, start=1):
            page_text = page.get_text("text").strip()
            if page_text:
                chunks.append(f"--- STRONA {index} ---\n{page_text}")
            else:
                chunks.append(f"--- STRONA {index} ---\n")

    text = "\n\n".join(chunks).strip()
    if not text or len(text.replace("-", "").strip()) < 20:
        return (
            "Nie udało się odczytać tekstu z PDF. Dokument może być skanem i wymagać OCR. "
            "TODO: dodać OCR w kolejnej wersji."
        )
    return text


def read_docx(file: BinaryIO) -> str:
    """Read paragraphs and tables from a DOCX file."""
    data = _read_bytes(file)
    document = Document(BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"--- TABELA {table_index} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()


def read_txt(file: BinaryIO) -> str:
    """Read plain text, preferring UTF-8 and falling back to latin-1."""
    data = _read_bytes(file)
    try:
        return data.decode("utf-8").strip()
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace").strip()


def read_catalog_file(uploaded_file: BinaryIO) -> str:
    """Dispatch catalog reading by uploaded file extension."""
    name = getattr(uploaded_file, "name", "")
    suffix = Path(name).suffix.lower()

    if suffix == ".pdf":
        return read_pdf(uploaded_file)
    if suffix == ".docx":
        return read_docx(uploaded_file)
    if suffix == ".txt":
        return read_txt(uploaded_file)

    raise ValueError("Nieobsługiwany format karty katalogowej. Dozwolone: PDF, DOCX, TXT.")
